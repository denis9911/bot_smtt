from telebot import types
from create_bot import bot, ScheduleStates
from datetime import datetime as dt
from openpyxl import load_workbook
import os
import docx

if os.path.exists(r'\\192.168.0.5\pool1\user\Миллер К.М\tgbot\raspisanie.xlsx'):
    lw = load_workbook(filename=r'\\192.168.0.5\pool1\user\Миллер К.М\tgbot\raspisanie.xlsx')
    sheet_full = lw.active  # Выбрать активный лист


@bot.message_handler(state=ScheduleStates.full_list_changes)
def full_list_changes(message):
    path = r'\\192.168.0.5\pool1\user\Миллер К.М\tgbot'
    files = []
    word_from_doc = []
    word_final = []
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add(types.KeyboardButton('Вернуться в главное меню'))
    for folder in os.listdir(path):
        if message.text.lower() in folder:
            files.append(os.path.join(path, folder))
    if len(files) < 1:
        if message.text == 'Вернуться в главное меню':
            bot.delete_state(message.from_user.id, message.chat.id)
            return bot.send_message(message.chat.id, text='Нажмите на кнопку ещё раз', reply_markup=markup)
        else:
            return bot.send_message(message.chat.id,
                                    text='Неправильная команда, нажмите кнопку выбора из представленных вариантов...')
    doc_file = docx.Document(files[0])
    for tables in doc_file.tables:
        for row in tables.rows:
            for cell in row.cells:
                word_from_doc.append(cell.text)
    for word_index in range(0, len(word_from_doc), 4):
        word_final.append('\n'.join(word_from_doc[
                                    word_index:word_index + 4]) + f'\n{"-" * 30}\n')  # Объединяем первые 4 элемента списка в 1, в конце добавляем подчеркивание
    bot.send_message(message.chat.id, ''.join(word_final))


@bot.message_handler(state=ScheduleStates.week_in_a_row)
def week_in_a_row(message):
    week_number = dt.strftime(dt.now(), '%W')

    def parity_week_next():
        markup = types.ReplyKeyboardMarkup(row_width=2)
        btn1 = types.KeyboardButton('Понедельник')
        btn2 = types.KeyboardButton('Вторник')
        btn3 = types.KeyboardButton('Среда')
        btn4 = types.KeyboardButton('Четверг')
        btn5 = types.KeyboardButton('Пятница')
        btn6 = types.KeyboardButton('Суббота')
        markup.add(btn1, btn2, btn3, btn4, btn5, btn6)
        bot.send_message(message.chat.id, text='Выберите день недели:', reply_markup=markup)
        bot.set_state(message.from_user.id, ScheduleStates.day_of_the_week, message.chat.id)
        with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
            data['parity_week'] = parity_week

    if message.text == 'Вернуться в главное меню':
        markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
        markup.add(types.KeyboardButton('Вернуться в главное меню'))
        bot.send_message(message.chat.id, text='Нажмите на кнопку ещё раз', reply_markup=markup)
        bot.delete_state(message.from_user.id, message.chat.id)
    elif message.text == 'Следующая неделя' or message.text == 'Текущая неделя':
        if message.text == 'Следующая неделя':
            with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
                week_number = int(week_number) + 1
                data['week_number'] = week_number
        else:
            with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
                week_number = int(week_number)
                data['week_number'] = week_number
        if int(week_number) % 2 == 0:
            parity_week = 'знаменатель'
        else:
            parity_week = 'числитель'
        parity_week_next()
    else:
        bot.send_message(message.chat.id,
                         text='Неправильная команда, нажмите кнопку выбора из представленных вариантов...')


@bot.message_handler(state=ScheduleStates.day_of_the_week)
def day_of_the_week(message):
    week_count = []

    def day_week_next():
        markup = types.ReplyKeyboardMarkup(row_width=2)
        btn1 = types.KeyboardButton('1 курс')
        btn2 = types.KeyboardButton('2 курс')
        btn3 = types.KeyboardButton('3 курс')
        btn4 = types.KeyboardButton('4 курс')
        markup.add(btn1, btn2, btn3, btn4)
        bot.send_message(message.chat.id, text='Выберите курс:', reply_markup=markup)
        bot.set_state(message.from_user.id, ScheduleStates.group_select, message.chat.id)
        with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
            data['week_count'] = week_count
            data['user_day'] = message.text

    def append_week(integer1, integer2):
        week_count.append(integer1), week_count.append(integer2)

    if message.text == 'Понедельник':
        append_week(9, 17)
        day_number(message, 1)
        day_week_next()
    elif message.text == 'Вторник':
        append_week(18, 26)
        day_number(message, 2)
        day_week_next()
    elif message.text == 'Среда':
        append_week(28, 36)
        day_number(message, 3)
        day_week_next()
    elif message.text == 'Четверг':
        append_week(37, 45)
        day_number(message, 4)
        day_week_next()
    elif message.text == 'Пятница':
        append_week(46, 54)
        day_number(message, 5)
        day_week_next()
    elif message.text == 'Суббота':
        append_week(55, 61)
        day_number(message, 6)
        day_week_next()
    else:
        bot.send_message(message.chat.id,
                         text='Неправильная команда, нажмите кнопку выбора из представленных вариантов...')


@bot.message_handler(state=ScheduleStates.group_select)
def group_select(message):
    bot.send_message(message.chat.id, text='Загружаю список групп...')
    groups = []
    buttons = []
    # Ищем в строке все значения с нужным курсом, для добавления в кнопки
    if message.text[0] == '1' or message.text[0] == '2' or message.text[0] == '3' or message.text[0] == '4':
        for rows in sheet_full.iter_rows(min_row=8, max_row=8):
            for cell in rows:
                if cell.value is not None:
                    if str(message.text[0]) in cell.value[0][0]:
                        groups.append(cell.value)
                        buttons.append(types.KeyboardButton(cell.value))  # Получили зачения групп нужного курса
        markup = types.ReplyKeyboardMarkup()
        markup.add(*buttons)
        bot.send_message(message.chat.id, text='Выберите нужную группу:', reply_markup=markup)
        bot.set_state(message.from_user.id, ScheduleStates.raspisanie, message.chat.id)
        with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
            data['groups'] = groups
    else:
        bot.send_message(message.chat.id,
                         text='Неправильная команда, нажмите кнопку выбора из представленных вариантов...')


@bot.message_handler(state=ScheduleStates.raspisanie)
def raspisanie(message):
    bot.send_message(message.chat.id, text='Загружаю расписание исходя из вашего выбора...')
    with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
        if message.text in data['groups']:
            data['user_group'] = message.text
            group_letter = ''
            for i in sheet_full.iter_rows(8, 8,
                                          max_col=50):  # Функция поиска по строке, ищем в 8 строке(где все группы) аргумент функции
                for j in i:
                    if message.text == j.value:
                        group_letter = j.column_letter  # Получили значение колонки нужного названия группы
            list_lents = []
            list_chislitel = []
            list_znamenatel = []
            count_chislitel = 1
            count_znamenatel = 1
            for i in range(data['week_count'][0], data['week_count'][1], 2):
                list_chislitel.append(sheet_full[group_letter + str(i)].value)
            for i in range(data['week_count'][0] + 1, data['week_count'][1], 2):
                list_znamenatel.append(sheet_full[group_letter + str(i)].value)
            # Если числитель, то перебираем значения, если значение отсуствует, то печатаем 🪟, если есть, то печатаем его
            if data['parity_week'] == 'числитель':
                for i in range(len(list_chislitel)):
                    if type(list_chislitel[i]) == str:
                        list_lents.append(
                            f'{count_chislitel} пара - {list_chislitel[i]}\n{"-" * 30}')  # Печатаем подряд значения из списка, если неделя по числителю
                        count_chislitel += 1  # Какая по счёту пара
                    else:
                        list_lents.append(
                            f'{count_chislitel} пара - ❌❌❌\n{"-" * 30}')  # Если None
                        count_chislitel += 1
            # Если нечетная, то перебираем значения, если первое NoneType, то печатаем первое из четного листа, если и там и там нет значения, то ❌❌❌
            if data['parity_week'] == 'знаменатель':
                for i in range(len(list_znamenatel)):
                    # если значение элемента строки не NoneType
                    if type(list_znamenatel[i]) == str:
                        list_lents.append(f'{count_znamenatel} пара - {list_znamenatel[i]}\n{"-" * 30}')
                        count_znamenatel += 1
                    # если значение элемента строки знаменателя NoneType, но в числителе есть пары, то печатаем числитель
                    elif type(list_znamenatel[i]) != str and type(list_chislitel[i]) == str:
                        list_lents.append(f'{count_znamenatel} пара - {list_chislitel[i]}\n{"-" * 30}')
                        count_znamenatel += 1
                    # если значения не NoneType отсутствуют - пары нет
                    else:
                        list_lents.append(f'{count_znamenatel} пара - ❌❌❌\n{"-" * 30}')
                        count_znamenatel += 1
            izmenenia(message)
            user_selected_date_raw = f'{dt.now().strftime("%Y")}-{data["week_number"]}-{data["day_number"]}'
            user_selected_date = dt.strptime(user_selected_date_raw, '%Y-%U-%w')
            perenos = '\n'  # f-строки не могут включать в себя обратный слеш
            bot.send_message(message.chat.id,
                             text=f'\n{data["user_group"]}, {user_selected_date.strftime("%d.%m.%Y")} ({data["user_day"]}), {data["parity_week"].title()}\n{"-" * 30}\n{perenos.join(list_lents)}')
            markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
            markup.add(types.KeyboardButton('Вернуться в главное меню'))
            bot.send_message(message.chat.id, text='Можете вернуться в главное меню:', reply_markup=markup)
        else:
            return bot.send_message(message.chat.id,
                                    text='Неправильная команда, нажмите кнопку выбора из представленных вариантов...')
    bot.delete_state(message.from_user.id, message.chat.id)


# утил. функция для изменений
def izmenenia(message):
    path = r'\\192.168.0.5\pool1\user\Миллер К.М\tgbot'  # Отправляем папку
    files = []
    word = []
    user_group_rows = []
    for root, d, folder in os.walk(path):
        for doc in folder:
            with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
                if data['user_day'].lower() in doc:
                    files.append(os.path.join(root, doc))
    if len(files) > 0:
        doc_file = docx.Document(files[0])
        for tables in doc_file.tables:
            for row in tables.rows:
                for cell in row.cells:
                    word.append(cell.text)
            for i in range(len(word)):
                if message.text[:6] in word[i]:
                    user_group_rows.append(word[i - 1:i + 3])
        if len(user_group_rows) > 0:
            for rows in user_group_rows:
                new_line_char = '\n'
                bot.send_message(message.chat.id,
                                 text=f'*Изменения в расписании на этот день недели:\n{new_line_char.join(rows)}*',
                                 parse_mode='Markdown')
        else:
            return bot.send_message(message.chat.id, text='*Изменений в расписании для выбранного дня недели нет*',
                                    parse_mode='Markdown')
    else:
        return bot.send_message(message.chat.id, text='*Изменений в расписании для выбранного дня недели нет*',
                                parse_mode='Markdown')


def day_number(message, num):
    with bot.retrieve_data(message.from_user.id, message.chat.id) as data:
        data['day_number'] = num
